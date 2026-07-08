import io
import re
import json
import string
import urllib.request
import urllib.parse
import datetime
import html as html_lib
import requests
from docx import Document
from collections import defaultdict
import fire
import pypdf


def handle_web_fallback(url):
    try:
        req_url = url
        if url.lower().endswith('.pdf') and ('biorxiv.org' in url.lower() or 'medrxiv.org' in url.lower()):
            req_url = re.sub(r'.full.pdf$', '', url, flags=re.IGNORECASE).replace('.pdf', '')
        req = urllib.request.Request(req_url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9'
        })
        with urllib.request.urlopen(req, timeout=5) as response:
            content_type = response.headers.get('Content-Type', '').lower()
            if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
                try:
                    pdf_bytes = response.read()
                    with pypdf.PdfReader(io.BytesIO(pdf_bytes)) as reader:
                        meta = reader.metadata
                        title = meta.title.strip() if (meta and meta.title) else ""
                        author = [meta.author.strip()] if (meta and meta.author) else ["Unknown"]
                        if title and title.lower() not in ['blocked', '403 forbidden', 'not found']:
                            return author, "2026", title, "Journal PDF"
                except Exception: pass
                html_content = ""
            else:
                html_content = response.read().decode('utf-8', errors='ignore')

            title_match = re.search(r'<meta\s+name="citation_title"\s+content="(.*?)"', html_content, re.IGNORECASE) or \
                          re.search(r'<meta\s+property="og:title"\s+content="(.*?)"', html_content, re.IGNORECASE) or \
                          re.search(r'<title>(.*?)</title>', html_content, re.DOTALL | re.IGNORECASE)
            title = title_match.group(1).strip() if title_match else ""
            title = re.sub(r'\s+', ' ', title)

            CHALLENGE_KEYWORDS = ['verifying your browser', 'cloudflare', 'attention required', 'robot check', 'captcha', '403 forbidden']
            if any(ck in title.lower() for ck in CHALLENGE_KEYWORDS): title = ""

            if title and title.lower() not in ['blocked', '403 forbidden', 'not found', 'error']:
                author_matches = re.findall(r'<meta\s+name="citation_author"\s+content="(.*?)"', html_content, re.IGNORECASE)
                if not author_matches:
                    fallback_author = re.search(r'<meta\s+name="author"\s+content="(.*?)"', html_content, re.IGNORECASE)
                    author_matches = [fallback_author.group(1)] if fallback_author else []
                authors = [au.strip() for au in author_matches if au.strip()]
                if not authors: authors = ["Unknown"]

                year_match = re.search(r'<meta\s+name="citation_publication_date"\s+content="(\d{4})"', html_content, re.IGNORECASE) or \
                             re.search(r'<meta\s+name="citation_date"\s+content="(\d{4})"', html_content, re.IGNORECASE) or \
                             re.search(r'\b(19\d{2}|20\d{2})\b', html_content)
                year = year_match.group(1) if year_match else "2026"

                journal_match = re.search(r'<meta\s+name="citation_journal_title"\s+content="(.*?)"', html_content, re.IGNORECASE) or \
                                re.search(r'<meta\s+property="og:site_name"\s+content="(.*?)"', html_content, re.IGNORECASE)
                journal = journal_match.group(1).strip() if journal_match else ""
                if not journal:
                    domain_match = re.search(r'https?://(?:www\.)?([^/]+)', url)
                    journal = domain_match.group(1).capitalize() if domain_match else "Web Resource"
                return [html_lib.unescape(a) for a in authors], year, html_lib.unescape(title), html_lib.unescape(journal)
    except Exception: pass

    if 'openreview.net' in url.lower():
        or_id = url.split('id=')[-1].split('&')[0]
        return ["OpenReview"], "2026", f"[OpenReview] {or_id}", "OpenReview Platform"
    domain_match = re.search(r'https?://(?:www\.)?([^/]+)', url)
    domain = domain_match.group(1).split('.')[0].capitalize() if domain_match else "Web"
    filename = url.split('/')[-1].split('?')[0].replace('.html', '').replace('.htm', '').replace('.pdf', '')
    title = f"[{domain}] {re.sub(r'[_+\-]+', ' ', filename).strip().capitalize()}"
    return [domain], "2026", title, "Web Resource"


def get_live_metadata(url):
    filename_author, filename_year, filename_title = "Unknown", "2026", ""
    NCBI_API_KEY = "af5c3517cfb808bc9c5809e468d9cc587708"
    ncbi_suffix = f"&api_key={NCBI_API_KEY}" if NCBI_API_KEY else ""
    try:
        ADMIN_DOMAINS = ['europa.eu', 'youtube.com', 'scribbr.com', 'grammar.com', 'languagetool.org', 'reescribirtextos.net', 'ahrefs.com', 'bio-techne.com', 'horizonteeuropa.es']
        if any(dom in url.lower() for dom in ADMIN_DOMAINS): return handle_web_fallback(url)

        doi, doi_query = "", ""
        doi_match = re.search(r'(10\.\d{4,9}/[^\s,)\"\]\?]+)', url)
        if doi_match:
            doi = doi_match.group(1)
            for suffix in ['.full.pdf', '.full-text', '.full', '.pdf', '/full', '/pdf', '/abstract']:
                if doi.lower().endswith(suffix): doi = doi[:-len(suffix)]
            doi = re.sub(r'v\d+$', '', doi, flags=re.IGNORECASE)

        if not doi:
            nature_match = re.search(r'nature\.com/articles/([a-zA-Z0-9.\-_]+)', url, re.IGNORECASE)
            if nature_match: doi = f"10.1038/{nature_match.group(1)}"

        if not doi:
            elife_match = re.search(r'elifesciences\.org/(?:articles|reviewed-preprints)/(\d+)', url, re.IGNORECASE)
            if elife_match: doi = f"10.7554/eLife.{elife_match.group(1)}"

            openreview_match = re.search(r'openreview\.net/(?:forum|pdf)\?id=([\w\-]+)', url, re.IGNORECASE)
            if openreview_match:
                or_id = openreview_match.group(1)
                for api_version in ['api2', 'api']:
                    try:
                        or_api = f"https://{api_version}.openreview.net/notes?id={or_id}"
                        req = urllib.request.Request(or_api, headers={'User-Agent': 'Mozilla/5.0'})
                        with urllib.request.urlopen(req, timeout=4) as response:
                            or_data = json.loads(response.read().decode())
                            if or_data.get('notes'):
                                note = or_data['notes'][0]
                                content = note.get('content', {})
                                t_val = content.get('title', {})
                                title = t_val.get('value', 'Document') if isinstance(t_val, dict) else t_val
                                a_val = content.get('authors', {})
                                authors = a_val.get('value', ['OpenReview']) if isinstance(a_val, dict) else a_val
                                if not isinstance(authors, list): authors = [authors]
                                timestamp = note.get('cdate') or note.get('tcdate') or 1767222000000
                                year = str(datetime.datetime.fromtimestamp(timestamp / 1000).year)
                                return authors, year, title, "OpenReview"
                    except Exception: pass

            arxiv_match = re.search(r'arxiv\.org/(?:pdf|abs)/(\d+\.\d+)', url, re.IGNORECASE)
            if arxiv_match:
                arxiv_id = arxiv_match.group(1)
                try:
                    arxiv_api = f"http://export.arxiv.org/api/query?id_list={arxiv_id}"
                    req = urllib.request.Request(arxiv_api, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=4) as response:
                        xml_data = response.read().decode()
                        entry_match = re.search(r'<entry>.*?</entry>', xml_data, re.DOTALL)
                        if entry_match:
                            entry = entry_match.group(0)
                            t_m = re.search(r'<title>(.*?)</title>', entry, re.DOTALL)
                            title = re.sub(r'\s+', ' ', t_m.group(1)).strip() if t_m else "arXiv Entry"
                            p_m = re.search(r'<published>(\d{4})', entry)
                            year = p_m.group(1) if p_m else "2026"
                            authors = re.findall(r'<name>(.*?)</name>', entry)
                            return (authors if authors else ["arXiv"]), year, title, "arXiv"
                except Exception: pass
                doi = f"10.48550/arXiv.{arxiv_id}"

            pmc_match = re.search(r'pmc\.ncbi\.nlm\.nih\.gov/articles/PMC(\d+)', url, re.IGNORECASE) or \
                        re.search(r'ncbi\.nlm\.nih\.gov/pmc/articles/PMC(\d+)', url, re.IGNORECASE) or \
                        re.search(r'[\/_]pmc(\d{5,8})', url, re.IGNORECASE)
            if pmc_match:
                pmcid = pmc_match.group(1)
                api_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?db=pmc&id={pmcid}&retmode=json{ncbi_suffix}"
                req = urllib.request.Request(api_url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=4) as response:
                    data = json.loads(response.read().decode())
                    info = data['result'][pmcid]
                    authors = [au['name'] for au in info.get('authors', [])] if info.get('authors') else ["Unknown"]
                    year_match = re.search(r'\b\d{4}\b', info.get('pubdate', '2026'))
                    year = year_match.group(0) if year_match else "2026"
                    title = info.get('title', 'Reference').strip('.')
                    return authors, year, title, info.get('source', 'PMC')

            pubmed_match = re.search(r'pubmed\.ncbi\.nlm\.nih\.gov/(\d+)', url, re.IGNORECASE)
            if pubmed_match:
                pmid = pubmed_match.group(1)
                api_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?db=pubmed&id={pmid}&retmode=json{ncbi_suffix}"
                req = urllib.request.Request(api_url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=4) as response:
                    data = json.loads(response.read().decode())
                    info = data['result'][pmid]
                    authors = [au['name'] for au in info.get('authors', [])] if info.get('authors') else ["Unknown"]
                    year_match = re.search(r'\b\d{4}\b', info.get('pubdate', '2026'))
                    year = year_match.group(0) if year_match else "2026"
                    title = info.get('title', 'Reference').strip('.')
                    return authors, year, title, info.get('source', 'PubMed')

        if not doi and url.lower().endswith('.pdf'):
            try:
                pdf_req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0', 'Accept': 'application/pdf'})
                with urllib.request.urlopen(pdf_req, timeout=6) as pdf_res:
                    pdf_bytes = pdf_res.read()
                    if pdf_bytes.startswith(b'%PDF'):
                        with pypdf.PdfReader(io.BytesIO(pdf_bytes)) as reader:
                            pdf_text = "".join(reader.pages[page_idx].extract_text() or "" for page_idx in range(min(2, len(reader.pages))))
                            pdf_doi_match = re.search(r'(10\.\d{4,9}/[^\s,)\"\]\?]+)', pdf_text)
                            if pdf_doi_match: doi = re.sub(r'v\d+$', '', pdf_doi_match.group(1), flags=re.IGNORECASE)
            except Exception: pass

        if not doi:
            pii_match = re.search(r'/pii/([\w\d]{14,18})', url, re.IGNORECASE)
            if pii_match: doi_query = f'"{pii_match.group(1)}"'
            elif 'escholarship.org' in url.lower():
                escholar_match = re.search(r'(?:item|content)/([\w\d]{8,10})', url)
                if escholar_match: doi_query = f"escholarship {escholar_match.group(1)}"
            elif 'academic.oup.com' in url.lower():
                oup_match = re.search(r'/article/(\d+)', url) or re.search(r'/article/[^/]+/(\d+)', url)
                if oup_match: doi_query = f"Oxford OUP {oup_match.group(1)}"
            elif url.lower().endswith('.pdf'):
                filename = url.split('/')[-1].replace('.pdf', '')
                filename_title = re.sub(r'[_+\-]+', ' ', filename).strip()
                fn_match = re.search(r'^([A-Za-z]+)(\d{4})', filename)
                if fn_match: filename_author, filename_year = fn_match.group(1), fn_match.group(2)
                query_text = re.sub(r'[_+\-]+', ' ', re.sub(r'(?<!^)(?=[A-Z])', ' ', filename)).strip()
                if len(query_text) > 8: doi_query = query_text

        item = None
        polite_headers = {'User-Agent': 'AcademicMetadataParser/2.0 (mailto:gurkpeter@yahoo.com; ResilientBatchScript)', 'Accept': 'application/json'}
        if doi:
            try:
                api_url = f"https://api.crossref.org/works/{urllib.parse.quote(doi, safe='/')}"
                req = urllib.request.Request(api_url, headers=polite_headers)
                with urllib.request.urlopen(req, timeout=5) as response: item = json.loads(response.read().decode())['message']
            except Exception: doi_query = f'"{doi}"'
        if not item and doi_query:
            try:
                api_url = f"https://api.crossref.org/works?query={urllib.parse.quote(doi_query)}&rows=1"
                req = urllib.request.Request(api_url, headers=polite_headers)
                with urllib.request.urlopen(req, timeout=5) as response:
                    items = json.loads(response.read().decode())['message']['items']
                    if items: item = items[0]
            except Exception: pass

        if item:
            title = item['title'][0] if item.get('title') and isinstance(item['title'], list) and item['title'] else "Academic Reference"
            journal = ""
            if item.get('container-title'): journal = item['container-title'][0] if isinstance(item['container-title'], list) else str(item['container-title'])
            elif item.get('institution'): journal = item['institution'][0].get('name', '') if isinstance(item['institution'], list) else item['institution'].get('name', '')
            if not journal: journal = "Journal Resource"
            authors = [f"{au.get('family', '')}, {au.get('given', '')}".strip(', ') for au in item.get('author', []) if au.get('family')]
            if not authors: authors = ["Unknown"]
            year = "2026"
            for date_field in ['published-print', 'published-online', 'posted', 'created', 'issued']:
                if item.get(date_field) and item[date_field].get('date-parts') and item[date_field]['date-parts'][0]:
                    year = str(item[date_field]['date-parts'][0][0])
                    break
            return authors, year, title, journal

        if filename_author != "Unknown": return [filename_author], filename_year, filename_title, "Journal PDF"
        return handle_web_fallback(url)
    except Exception:
        if filename_author != "Unknown": return [filename_author], filename_year, filename_title, "Journal PDF"
        return handle_web_fallback(url)


def optimize_biorxiv_versions(ref_num_to_url):
    base_to_max_ver = {}
    for ref_num, url in ref_num_to_url.items():
        if 'biorxiv.org' in url.lower():
            match = re.match(r'(https?://(?:www\.)?biorxiv\.org/content/10\.1101/[\d\.]+)(v\d+)?', url)
            if match:
                base_url, ver_str = match.group(1), match.group(2) if match.group(2) else "v1"
                ver_num = int(ver_str[1:])
                if base_url not in base_to_max_ver or ver_num > base_to_max_ver[base_url]['ver_num']:
                    base_to_max_ver[base_url] = {'ver_num': ver_num, 'ver_str': ver_str, 'full_url': f"{base_url}{ver_str}" if match.group(2) else url}
    for ref_num, url in list(ref_num_to_url.items()):
        if 'biorxiv.org' in url.lower():
            match = re.match(r'(https?://(?:www\.)?biorxiv\.org/content/10\.1101/[\d\.]+)(v\d+)?', url)
            if match and match.group(1) in base_to_max_ver:
                latest_url = base_to_max_ver[match.group(1)]['full_url']
                if url != latest_url:
                    print(f"🔄 Remapping bioRxiv marker [{ref_num}] -> latest ({base_to_max_ver[match.group(1)]['ver_str']})")
                    ref_num_to_url[ref_num] = latest_url
    return ref_num_to_url


def resolve_duplicate_titles_and_registry(url_metadata_registry, ref_num_to_url):
    url_remap, groups = {}, {}
    for url, meta in url_metadata_registry.items():
        base_author = re.sub(r'_\d+$', '', meta['display'].lower().strip())
        groups.setdefault((base_author, meta['year']), []).append((url, meta))

    def get_significant_words(title_str): return set(re.findall(r'\b[a-z]{3,}\b', title_str.lower()))

    for (author, year), items in groups.items():
        if len(items) < 2: continue
        merged_urls = set()
        for i in range(len(items)):
            url_i, meta_i = items[i]
            if url_i in merged_urls: continue
            words_i = get_significant_words(meta_i['title'])
            if not words_i: continue
            for j in range(i + 1, len(items)):
                url_j, meta_j = items[j]
                if url_j in merged_urls: continue
                words_j = get_significant_words(meta_j['title'])
                if not words_j: continue
                intersection = words_i.intersection(words_j)
                smaller_size = min(len(words_i), len(words_j))
                overlap_ratio = len(intersection) / smaller_size if smaller_size > 0 else 0
                if overlap_ratio >= 0.70:
                    is_i_preprint = 'biorxiv.org' in url_i.lower() or 'medrxiv.org' in url_i.lower()
                    is_j_preprint = 'biorxiv.org' in url_j.lower() or 'medrxiv.org' in url_j.lower()
                    if is_i_preprint and not is_j_preprint: master_url, dup_url = url_j, url_i
                    elif not is_i_preprint and is_j_preprint: master_url, dup_url = url_i, url_j
                    else:
                        if not meta_i.get('journal') or meta_i['journal'].lower() in ['web resource', 'web', 'journal pdf', 'unknown']: master_url, dup_url = url_j, url_i
                        else: master_url, dup_url = url_i, url_j
                    print(f"⚠️ Duplicate Paper Detected: Merging {dup_url} -> {master_url}")
                    url_remap[dup_url] = master_url
                    merged_urls.add(dup_url)
                    for ref_num, current_url in list(ref_num_to_url.items()):
                        if current_url == dup_url: ref_num_to_url[ref_num] = master_url
                    if dup_url in url_metadata_registry: del url_metadata_registry[dup_url]
    return url_metadata_registry, ref_num_to_url, url_remap


def fetch_doi_from_meta_tags(url):
    try:
        res = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
        if res.status_code == 200:
            doi_match = re.search(r'<meta\s+name=["\']citation_doi["\']\s+content=["\']([^"\']+)["\']', res.text, re.IGNORECASE)
            if doi_match: return doi_match.group(1).strip()
    except Exception: pass
    return None


_LEADING_REF_NUM_RE = re.compile(r'^\s*[\[(]?(\d{1,4})[\])\-:.]\s+')


def _extract_leading_ref_number(text):
    match = _LEADING_REF_NUM_RE.match(text)
    return match.group(1) if match else None


def register_reference(ref_num_to_url, used_ref_nums, candidate_ref_num, url, collision_log):
    ref_num = candidate_ref_num
    if ref_num in used_ref_nums:
        original, suffix = ref_num, 1
        while ref_num in used_ref_nums:
            suffix += 1
            ref_num = f"{original}_dup{suffix}"
        collision_log.append((original, ref_num, ref_num_to_url.get(original), url))
    used_ref_nums.add(ref_num)
    ref_num_to_url[ref_num] = url
    return ref_num


def print_collision_summary(ref_num_collisions):
    print("\n📊 Collision Summary")
    if ref_num_collisions:
        print(f"⚠️  {len(ref_num_collisions)} numbering collision(s) found.")
    else:
        print("✅ No numbering collisions found.")


def normalize_author_name(raw_author):
    """Strictly normalizes to 'Lastname, Initials' for RIS/Zotero to ensure consistent author matching."""
    author = re.sub(r'\s+', ' ', raw_author.strip())
    if not author: return "Unknown"

    if ',' in author:
        parts = [p.strip() for p in author.split(',', 1)]
        lastname, given = parts[0], parts[1] if len(parts) > 1 else ""
    else:
        tokens = author.split()
        if len(tokens) == 1: return tokens[0]
        last_token = tokens[-1].replace('.', '')
        if len(last_token) <= 3 and last_token.isalpha() and last_token.isupper():
            lastname, given = tokens[0], " ".join(tokens[1:])
        else:
            lastname, given = tokens[-1], " ".join(tokens[:-1])

    given_parts = given.split()
    formatted_given = []
    for p in given_parts:
        clean_p = p.replace('.', '').strip()
        if not clean_p: continue

        # Handle blocks of initials like "GM" -> "G. M."
        if len(clean_p) > 1 and clean_p.isalpha() and clean_p.isupper():
            for char in clean_p:
                formatted_given.append(f"{char}.")
        else:
            # Handle full names or single initials -> "G."
            formatted_given.append(f"{clean_p[0]}.")

    given_str = " ".join(formatted_given)
    return f"{lastname}, {given_str}" if given_str else lastname


def get_display_author(clean_authors, fallback="Unknown"):
    if not clean_authors: return fallback
    first_author = clean_authors[0]
    if ',' in first_author: lastname = first_author.split(',')[0].strip()
    else:
        tokens = first_author.split()
        if not tokens: return fallback
        if len(tokens[-1].replace('.', '')) <= 3 and tokens[-1].replace('.', '').isalpha(): lastname = tokens[0]
        else: lastname = tokens[-1]
    lastname = re.sub(r'\s+[A-Za-z]\b.*$', '', lastname).strip().rstrip('.')
    return f"{lastname} et al." if len(clean_authors) > 1 else lastname


def extract_cited_numbers(text):
    """Scans document text for all cited reference numbers."""
    cited_nums = set()
    # Correct regex to match the whole bracket content
    brackets = re.findall(r'\[([\d,\-\s]+)\]', text)
    for bracket in brackets:
        parts = bracket.split(',')
        for part in parts:
            part = part.strip()
            if '-' in part:
                try:
                    start, end = part.split('-')
                    for i in range(int(start.strip()), int(end.strip()) + 1):
                        cited_nums.add(str(i))
                except ValueError:
                    pass
            else:
                if part.isdigit():
                    cited_nums.add(part)
    return cited_nums


def convert(input_docx, output_docx="marked_document.docx", output_ris="zotero_import.ris"):
    doc = Document(input_docx)
    ref_num_to_url, used_ref_nums, ref_num_collisions = {}, set(), []
    url_pattern = re.compile(r'https?://\S+')
    text_ref_pattern = re.compile(r'(\d+)\s*[.\s\t)\-:]+\s*(https?://\S+)')
    native_list_counter = 1

    print("🔍 Step 1: Scanning document structure for references...")
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        url_match = url_pattern.search(text)
        if not url_match: continue
        url = url_match.group(0).rstrip('.,] ) ')
        text_match = text_ref_pattern.search(text)
        if text_match: candidate_ref_num = text_match.group(1)
        elif bool(para._element.xpath('./w:pPr/w:numPr')) or para.style.name.startswith('List Number'):
            candidate_ref_num, native_list_counter = str(native_list_counter), native_list_counter + 1
        else:
            leading_num = _extract_leading_ref_number(text)
            if leading_num: candidate_ref_num = leading_num
            else: candidate_ref_num, native_list_counter = str(native_list_counter), native_list_counter + 1
        register_reference(ref_num_to_url, used_ref_nums, candidate_ref_num, url, ref_num_collisions)

    if not ref_num_to_url:
        print("❌ Error: Could not extract references.")
        return

    # 🔴 CROSS-CHECK FILTER
    full_text = "\n".join([p.text for p in doc.paragraphs])
    cited_nums = extract_cited_numbers(full_text)
    original_count = len(ref_num_to_url)
    ref_num_to_url = {k: v for k, v in ref_num_to_url.items() if k in cited_nums}
    print(f"✂️  Filtered References: Kept {len(ref_num_to_url)} cited URLs out of {original_count} total listed URLs.")
    if not ref_num_to_url: return

    ref_num_to_url = optimize_biorxiv_versions(ref_num_to_url)
    unique_urls = list(dict.fromkeys(ref_num_to_url.values()))

    # 🔴 STEP 2: PYTHONIC O(1) CITEKEY GENERATION
    print("\n🌐 Step 2: Fetching deep metadata...")
    url_metadata_registry = {}
    citekey_counts = defaultdict(int)
    SUFFIXES = [''] + list(string.ascii_lowercase)

    for url in unique_urls:
        meta = None
        if 'academic.oup.com' in url.lower():
            oup_match = re.search(r'/article/.*?/(\d+)(?:\?|$)', url) or re.search(r'/article/(\d+)', url)
            if oup_match:
                scraped_doi = fetch_doi_from_meta_tags(url)
                if scraped_doi: meta = fetch_doi_from_meta_tags and __import__('sys').modules[__name__].fetch_crossref_by_doi(scraped_doi)
                if not meta: meta = query_crossref_by_text(f"Oxford OUP {oup_match.group(1)}")
        if not meta: meta = get_live_metadata(url)

        if meta:
            authors, year, title, journal = meta
            authors = [normalize_author_name(au) for au in authors if au.strip()]
            if not authors: authors = ["Unknown"]
            display_name = get_display_author(authors)
            title = re.sub(r'\s+', ' ', title).strip()
            base_author = display_name.replace(' et al.', '').strip()
            base_citekey = f"{re.sub(r'[^a-zA-Z]', '', base_author)}{year}"

            count = citekey_counts[base_citekey]
            suffix = SUFFIXES[count] if count < len(SUFFIXES) else str(count)
            citekey = f"{base_citekey}{suffix}"
            citekey_counts[base_citekey] += 1

            url_metadata_registry[url] = {
                'authors': authors, 'display': display_name, 'year': year,
                'title': title, 'journal': journal, 'citekey': citekey,
                'disambiguation_suffix': suffix
            }
            print(f"  ✅ Asset Found: {display_name} ({year}{suffix})")

    url_metadata_registry, ref_num_to_url, _ = resolve_duplicate_titles_and_registry(url_metadata_registry, ref_num_to_url)

    # 🔴 STEP 3: Robust Multi-Citation Parser
    print("\n📝 Step 3: Upgrading document in-text citation markers...")

    def replace_bracket(match):
        inner = match.group(1)
        keys = []
        parts = inner.split(',')
        for part in parts:
            part = part.strip()
            if '-' in part:
                try:
                    start, end = part.split('-')
                    for i in range(int(start.strip()), int(end.strip()) + 1):
                        num = str(i)
                        if num in ref_num_to_url and ref_num_to_url[num] in url_metadata_registry:
                            keys.append(f"@{url_metadata_registry[ref_num_to_url[num]]['citekey']}")
                except ValueError:
                    pass
            else:
                if part.isdigit() and part in ref_num_to_url:
                    url = ref_num_to_url[part]
                    if url in url_metadata_registry:
                        keys.append(f"@{url_metadata_registry[url]['citekey']}")
        unique_keys = list(dict.fromkeys(keys))
        if unique_keys:
            return "[" + "; ".join(unique_keys) + "]"
        return match.group(0)

    def merge_adjacent_pandoc(match):
        """Merges adjacent Pandoc citations like [@A], [@B] or [@A] [@B] into a single multi-citation [@A; @B]"""
        content = match.group(0)
        keys = re.findall(r'(@[^\]\s;]+)', content)
        unique_keys = list(dict.fromkeys(keys))
        return "[" + "; ".join(unique_keys) + "]"

    for para in doc.paragraphs:
        if para.text:
            # 1. Match both square [] and round () brackets for citations
            para.text = re.sub(r'[\[\(]([\d,\-\s]+)[\]\)]', replace_bracket, para.text)
            # 2. Merge adjacent generated Pandoc citations to ensure ODF scan yields a single multi-citation field
            para.text = re.sub(r'(\[@[^\]]+\](?:\s*[,;]?\s*(?:and)?\s*\[@[^\]]+\])+)', merge_adjacent_pandoc, para.text)

    doc.save(output_docx)
    print(f"💾 File updates successfully written to: {output_docx}")

    # 🔴 STEP 4: CLEAN RIS EXPORT
    output_ris_file = output_ris if output_ris.endswith('.ris') else output_ris.replace('.bib', '.ris')
    print(f"\n📁 Step 4: Formatting and outputting Zotero RIS database to {output_ris_file}...")
    with open(output_ris_file, 'w', encoding='utf-8') as ris_file:
        for url, data in url_metadata_registry.items():
            citekey = data['citekey']
            pure_year = data['year']  # NO SUFFIX IN PY TAG!
            ris_file.write("TY  - JOUR\n")
            ris_file.write(f"ID  - {citekey}\n")
            ris_file.write(f"TI  - {data['title']}\n")
            for author in data['authors']:
                ris_file.write(f"AU  - {author}\n")
            ris_file.write(f"PY  - {pure_year}\n")
            if data.get('journal'):
                ris_file.write(f"JO  - {data['journal']}\n")
            ris_file.write(f"UR  - {url}\n")
            ris_file.write(f"M2  - Citation Key: {citekey}\n")
            ris_file.write("ER  - \n\n")

    print(f"💾 RIS file successfully written.")
    print_collision_summary(ref_num_collisions)
    print(f"✅ Success!")


if __name__ == '__main__':
    fire.Fire(convert)